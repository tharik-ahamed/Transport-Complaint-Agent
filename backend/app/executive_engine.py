"""
Phase 6 Executive Intelligence Engine
======================================
Implements:
  - handle_copilot_query()        — AI Copilot (OpenAI → Gemini → Local fallback)
  - compute_health_index()         — Transport Health Index
  - get_governance_recommendations() — Smart governance recommendations
  - get_heatmap_data()             — Geographic complaint hotspots
  - get_ai_explanation()           — Explainable AI per complaint
  - generate_report_data()         — Aggregate metrics for reports
  - build_pdf_report()             — ReportLab PDF bytes
  - build_docx_report()            — python-docx DOCX bytes
"""
from __future__ import annotations

import io
import json
import logging
import os
from collections import Counter, defaultdict
from datetime import datetime
from typing import Any, Optional

from sqlalchemy.orm import Session
from sqlalchemy import text

from app.models import Complaint
from app.predictive_models import RouteRisk, DriverRisk, BusRisk
from app.config import AI_ENABLED, GEMINI_API_KEY, GEMINI_MODEL

logger = logging.getLogger(__name__)

# ── Config: optional API keys ─────────────────────────────────────────
OPENAI_API_KEY: str = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL: str = os.getenv("OPENAI_MODEL", "gpt-4o-mini")


# ── Helpers ───────────────────────────────────────────────────────────
def _parse_json(value: Optional[str], fallback: list) -> list:
    if not value:
        return fallback
    try:
        return json.loads(value)
    except Exception:
        return fallback


# ══════════════════════════════════════════════════════════════════════
# COPILOT: Rule-based local fallback engine
# ══════════════════════════════════════════════════════════════════════

def _local_copilot(db: Session, question: str) -> dict:
    """
    Rule-based fallback — answers common operational questions without any
    external API.  Returns {"answer": str, "supporting_data": dict}.
    """
    q = question.lower()

    # ── Route complaints ───────────────────────────────────────────────
    if "route" in q and ("highest" in q or "most" in q or "complaint" in q or "problem" in q):
        rows = db.execute(
            text("SELECT route_number, COUNT(*) as cnt FROM complaints GROUP BY route_number ORDER BY cnt DESC LIMIT 10")
        ).fetchall()
        supporting = [{"route": r[0], "count": r[1]} for r in rows]
        if rows:
            top = rows[0]
            ans = (
                f"Route {top[0]} has the highest number of complaints with {top[1]} reports recorded. "
                f"Top routes: {', '.join(f'Route {r[0]} ({r[1]})' for r in rows[:3])}."
            )
        else:
            ans = "No route complaint data is available yet."
        return {"answer": ans, "supporting_data": {"top_routes": supporting}}

    # ── Critical / Safety ──────────────────────────────────────────────
    if "critical" in q or "safety" in q or "accident" in q or "emergency" in q:
        rows = db.execute(
            text(
                "SELECT complaint_id, passenger_name, incident_location, severity, category "
                "FROM complaints WHERE severity='Critical' OR category='Safety Issue' "
                "ORDER BY created_at DESC LIMIT 10"
            )
        ).fetchall()
        supporting = [
            {"id": r[0], "passenger": r[1], "location": r[2], "severity": r[3], "category": r[4]}
            for r in rows
        ]
        ans = (
            f"Found {len(rows)} critical/safety complaints. "
            + (f"Most recent: {rows[0][2]} ({rows[0][3]})." if rows else "None recorded.")
        )
        return {"answer": ans, "supporting_data": {"critical_complaints": supporting}}

    # ── Driver / Risk ──────────────────────────────────────────────────
    if "driver" in q or "operator" in q or "risk score" in q or ("risk" in q and "driver" in q):
        try:
            drivers = db.query(DriverRisk).order_by(DriverRisk.risk_score.desc()).limit(10).all()
            supporting = [
                {"bus": d.driver_identifier, "risk_score": round(d.risk_score, 2),
                 "risk_level": d.risk_level, "complaints": d.complaint_count}
                for d in drivers
            ]
            if drivers:
                top = drivers[0]
                ans = (
                    f"Driver/bus proxy {top.driver_identifier} has the highest risk score of "
                    f"{top.risk_score:.2f} ({top.risk_level}) with {top.complaint_count} complaints. "
                    f"Total high-risk operators: {sum(1 for d in drivers if d.risk_level in ('Critical','High'))}."
                )
            else:
                ans = "No driver risk assessments are available. Run the predictive analytics to generate them."
            return {"answer": ans, "supporting_data": {"driver_risks": supporting}}
        except Exception:
            pass

    # ── Retraining ─────────────────────────────────────────────────────
    if "retrain" in q or "training" in q:
        try:
            drivers = db.query(DriverRisk).filter(
                DriverRisk.risk_level.in_(["Critical", "High"])
            ).order_by(DriverRisk.risk_score.desc()).all()
            supporting = [
                {"bus": d.driver_identifier, "risk_level": d.risk_level,
                 "misconduct": d.misconduct_count}
                for d in drivers
            ]
            ans = (
                f"{len(drivers)} driver(s) require retraining based on risk assessment. "
                + (f"Immediate attention: {drivers[0].driver_identifier} (score {drivers[0].risk_score:.2f})." if drivers else "")
            )
            return {"answer": ans, "supporting_data": {"drivers_needing_retraining": supporting}}
        except Exception:
            pass

    # ── Maintenance / Buses ────────────────────────────────────────────
    if "maintenance" in q or "bus health" in q or ("bus" in q and ("repair" in q or "breakdown" in q)):
        try:
            buses = db.query(BusRisk).order_by(BusRisk.risk_score.desc()).limit(10).all()
            supporting = [
                {"bus": b.bus_number, "maintenance_count": b.maintenance_count,
                 "risk": b.maintenance_risk}
                for b in buses
            ]
            ans = (
                f"{len(buses)} buses have maintenance risk flags. "
                + (f"Bus {buses[0].bus_number} requires immediate inspection ({buses[0].maintenance_count} reports)." if buses else "")
            )
            return {"answer": ans, "supporting_data": {"buses_needing_maintenance": supporting}}
        except Exception:
            pass

    # ── Delays ────────────────────────────────────────────────────────
    if "delay" in q or "late" in q or "on time" in q:
        rows = db.execute(
            text(
                "SELECT incident_location, COUNT(*) as cnt FROM complaints "
                "WHERE category='Bus Delay' GROUP BY incident_location ORDER BY cnt DESC LIMIT 10"
            )
        ).fetchall()
        total_delays = db.execute(
            text("SELECT COUNT(*) FROM complaints WHERE category='Bus Delay'")
        ).scalar() or 0
        supporting = [{"location": r[0], "count": r[1]} for r in rows]
        top_locs = ", ".join(r[0] for r in rows[:3]) if rows else "N/A"
        ans = (
            f"Found {total_delays} delay-related complaints. "
            f"Top delay hotspots: {top_locs}."
        )
        return {"answer": ans, "supporting_data": {"delay_hotspots": supporting}}

    # ── More buses / frequency ─────────────────────────────────────────
    if "more buses" in q or "add buses" in q or "frequency" in q or "overcrowd" in q:
        rows = db.execute(
            text(
                "SELECT route_number, COUNT(*) as cnt FROM complaints "
                "WHERE category IN ('Overcrowding','Bus Delay') "
                "GROUP BY route_number ORDER BY cnt DESC LIMIT 5"
            )
        ).fetchall()
        supporting = [{"route": r[0], "overcrowding_delay_complaints": r[1]} for r in rows]
        ans = (
            "Routes most in need of additional buses based on overcrowding/delay complaints: "
            + (", ".join(f"Route {r[0]} ({r[1]} complaints)" for r in rows[:3]) if rows else "No data available.")
        )
        return {"answer": ans, "supporting_data": {"routes_needing_buses": supporting}}

    # ── Operational risks ─────────────────────────────────────────────
    if "operational risk" in q or "biggest risk" in q or "top risk" in q:
        try:
            high_routes = db.query(RouteRisk).filter(
                RouteRisk.risk_level.in_(["Critical", "High"])
            ).count()
            high_drivers = db.query(DriverRisk).filter(
                DriverRisk.risk_level.in_(["Critical", "High"])
            ).count()
            high_buses = db.query(BusRisk).filter(
                BusRisk.maintenance_risk.in_(["Critical", "High"])
            ).count()
            sla_breaches = db.execute(
                text("SELECT COUNT(*) FROM complaints WHERE sla_status='SLA Breached'")
            ).scalar() or 0
            ans = (
                f"Current operational risks: {high_routes} high-risk routes, "
                f"{high_drivers} high-risk drivers/operators, {high_buses} buses needing maintenance, "
                f"and {sla_breaches} SLA breaches. Prioritise route capacity and driver retraining."
            )
            return {
                "answer": ans,
                "supporting_data": {
                    "high_risk_routes": high_routes,
                    "high_risk_drivers": high_drivers,
                    "buses_needing_maintenance": high_buses,
                    "sla_breaches": sla_breaches,
                }
            }
        except Exception:
            pass

    # ── Generic fallback ───────────────────────────────────────────────
    total = db.execute(text("SELECT COUNT(*) FROM complaints")).scalar() or 0
    pending = db.execute(
        text("SELECT COUNT(*) FROM complaints WHERE complaint_status NOT IN ('Resolved','Closed')")
    ).scalar() or 0
    critical = db.execute(
        text("SELECT COUNT(*) FROM complaints WHERE severity='Critical'")
    ).scalar() or 0
    ans = (
        f"The transport system currently has {total} registered complaints, "
        f"{pending} pending resolution, and {critical} critical severity cases. "
        f"Ask about routes, drivers, delays, maintenance, or operational risks for detailed insights."
    )
    return {
        "answer": ans,
        "supporting_data": {"total": total, "pending": pending, "critical": critical}
    }


# ══════════════════════════════════════════════════════════════════════
# COPILOT: Main entry point — OpenAI → Gemini → Local
# ══════════════════════════════════════════════════════════════════════

def handle_copilot_query(db: Session, question: str) -> dict:
    """
    Answers natural language operational queries.
    Priority: OpenAI → Gemini → Local rule-based fallback.
    Always returns {"success": True, "answer": str, "supporting_data": dict, "source": str}.
    """

    # ── 1. Try OpenAI ──────────────────────────────────────────────────
    if OPENAI_API_KEY:
        try:
            from openai import OpenAI
            client = OpenAI(api_key=OPENAI_API_KEY)

            # Gather quick stats to ground the AI answer
            stats = _gather_quick_stats(db)

            system_prompt = (
                "You are an executive AI assistant for a public transport authority. "
                "Answer questions concisely and professionally using the operational data provided. "
                "Focus on actionable insights."
            )
            user_msg = (
                f"Question: {question}\n\n"
                f"Current operational snapshot:\n{json.dumps(stats, indent=2)}"
            )
            response = client.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_msg},
                ],
                max_tokens=400,
                temperature=0.3,
            )
            answer = response.choices[0].message.content.strip()
            return {
                "success": True,
                "answer": answer,
                "supporting_data": stats,
                "source": "openai",
            }
        except Exception as e:
            logger.warning(f"OpenAI copilot failed: {e}. Trying Gemini.")

    # ── 2. Try Gemini ──────────────────────────────────────────────────
    if GEMINI_API_KEY:
        try:
            import google.generativeai as genai
            genai.configure(api_key=GEMINI_API_KEY)
            model = genai.GenerativeModel(GEMINI_MODEL)

            stats = _gather_quick_stats(db)
            prompt = (
                f"You are an executive AI assistant for a public transport authority.\n"
                f"Question: {question}\n\n"
                f"Operational snapshot:\n{json.dumps(stats, indent=2)}\n\n"
                "Answer concisely and professionally with actionable insights."
            )
            response = model.generate_content(prompt)
            return {
                "success": True,
                "answer": response.text.strip(),
                "supporting_data": stats,
                "source": "gemini",
            }
        except Exception as e:
            logger.warning(f"Gemini copilot failed: {e}. Falling back to local engine.")

    # ── 3. Local rule-based fallback ───────────────────────────────────
    result = _local_copilot(db, question)
    return {
        "success": True,
        "answer": result["answer"],
        "supporting_data": result.get("supporting_data", {}),
        "source": "local",
    }


def _gather_quick_stats(db: Session) -> dict:
    """Gathers a compact operational snapshot to pass to AI models as context."""
    try:
        total = db.execute(text("SELECT COUNT(*) FROM complaints")).scalar() or 0
        critical = db.execute(
            text("SELECT COUNT(*) FROM complaints WHERE severity='Critical'")
        ).scalar() or 0
        sla_breaches = db.execute(
            text("SELECT COUNT(*) FROM complaints WHERE sla_status='SLA Breached'")
        ).scalar() or 0
        top_routes = db.execute(
            text(
                "SELECT route_number, COUNT(*) as cnt FROM complaints "
                "GROUP BY route_number ORDER BY cnt DESC LIMIT 5"
            )
        ).fetchall()
        top_categories = db.execute(
            text(
                "SELECT category, COUNT(*) as cnt FROM complaints "
                "GROUP BY category ORDER BY cnt DESC LIMIT 5"
            )
        ).fetchall()
        return {
            "total_complaints": total,
            "critical_complaints": critical,
            "sla_breaches": sla_breaches,
            "top_routes": [{"route": r[0], "count": r[1]} for r in top_routes],
            "top_categories": [{"category": r[0], "count": r[1]} for r in top_categories],
        }
    except Exception:
        return {}


# ══════════════════════════════════════════════════════════════════════
# TRANSPORT HEALTH INDEX
# ══════════════════════════════════════════════════════════════════════

def compute_health_index(db: Session) -> dict:
    """Returns a weighted Transport Health Index (0–100) with qualitative rating."""
    complaints = db.query(Complaint).all()
    total = len(complaints)
    if total == 0:
        return {"transport_health_score": 100, "rating": "Excellent", "metrics": {}}

    sla_breaches = sum(1 for c in complaints if c.sla_status == "SLA Breached")
    safeties = sum(1 for c in complaints if c.category == "Safety Issue" or c.severity == "Critical")
    misconducts = sum(1 for c in complaints if c.category in ("Driver Misconduct", "Conductor Misconduct"))
    maintenance = sum(1 for c in complaints if c.category == "Maintenance Issue")

    sla_score = max(100 - (sla_breaches / total * 200), 0)
    safety_score = max(100 - (safeties / total * 300), 0)
    driver_score = max(100 - (misconducts / total * 150), 0)
    bus_score = max(100 - (maintenance / total * 150), 0)
    volume_penalty = min(total / 50 * 10, 20)

    raw = (sla_score * 0.3 + safety_score * 0.3 + driver_score * 0.2 + bus_score * 0.2) - volume_penalty
    score = int(max(min(raw, 100), 0))
    rating = (
        "Excellent" if score >= 85 else
        "Good" if score >= 70 else
        "Average" if score >= 50 else
        "Poor"
    )
    return {
        "transport_health_score": score,
        "rating": rating,
        "metrics": {
            "sla_performance": round(sla_score, 1),
            "safety_rating": round(safety_score, 1),
            "driver_performance": round(driver_score, 1),
            "bus_health": round(bus_score, 1),
            "total_incidents": total,
        },
    }


# ══════════════════════════════════════════════════════════════════════
# GOVERNANCE RECOMMENDATIONS
# ══════════════════════════════════════════════════════════════════════

def get_governance_recommendations(db: Session) -> list[dict]:
    """Returns prioritised strategic governance recommendations."""
    recs: list[dict] = []

    try:
        route_risks = db.query(RouteRisk).order_by(RouteRisk.risk_score.desc()).limit(3).all()
        for r in route_risks:
            if r.risk_score >= 0.5:
                recs.append({
                    "category": "Route Optimization",
                    "title": f"Expand Capacity on Route {r.route_number}",
                    "description": (
                        f"Route {r.route_number} has risk score {r.risk_score:.2f}. "
                        "Add 2 additional buses to reduce wait times and crowding."
                    ),
                    "priority": "Critical" if r.risk_score >= 0.75 else "High",
                    "estimated_impact": "15–20% reduction in wait-time complaints",
                })
    except Exception:
        pass

    try:
        driver_risks = db.query(DriverRisk).order_by(DriverRisk.risk_score.desc()).limit(3).all()
        for d in driver_risks:
            if d.risk_score >= 0.5:
                recs.append({
                    "category": "Staff Training",
                    "title": f"Mandatory Safety Training for Bus {d.driver_identifier}",
                    "description": (
                        f"Operator on bus {d.driver_identifier} has {d.complaint_count} misconduct/safety "
                        "complaints. Require immediate safe-driving course."
                    ),
                    "priority": d.risk_level if d.risk_level in ("Critical", "High", "Medium") else "High",
                    "estimated_impact": "30% reduction in misconduct complaints",
                })
    except Exception:
        pass

    try:
        bus_risks = db.query(BusRisk).order_by(BusRisk.risk_score.desc()).limit(3).all()
        for b in bus_risks:
            if b.risk_score >= 0.5:
                recs.append({
                    "category": "Maintenance",
                    "title": f"Preventive Maintenance Audit for Bus {b.bus_number}",
                    "description": (
                        f"Bus {b.bus_number} shows {b.maintenance_count} engine/mechanical failure reports. "
                        "Inspect suspension, brakes, and heating systems."
                    ),
                    "priority": "High",
                    "estimated_impact": "Fewer breakdown delays",
                })
    except Exception:
        pass

    if len(recs) < 3:
        recs.append({
            "category": "General Operations",
            "title": "Establish Continuous Feedback Loops",
            "description": "Standardise resolution reviews and track repeat complaint patterns monthly.",
            "priority": "Medium",
            "estimated_impact": "Enhanced passenger satisfaction",
        })

    priority_order = ["Critical", "High", "Medium", "Low"]
    return sorted(recs, key=lambda x: priority_order.index(x.get("priority", "Low")) if x.get("priority", "Low") in priority_order else 3)


# ══════════════════════════════════════════════════════════════════════
# GEOGRAPHIC HEATMAP
# ══════════════════════════════════════════════════════════════════════

def get_heatmap_data(db: Session) -> dict:
    """Returns complaint density hotspots by incident location."""
    complaints = db.query(Complaint).all()
    hotspots: dict[str, dict] = defaultdict(lambda: {"count": 0, "safety": 0, "delay": 0, "maintenance": 0})

    for c in complaints:
        loc = (c.incident_location or "Unknown Location").strip()
        hotspots[loc]["count"] += 1
        cats = _parse_json(c.ai_categories, [c.category] if c.category else [])
        for cat in cats:
            if cat == "Safety Issue":
                hotspots[loc]["safety"] += 1
            elif cat == "Bus Delay":
                hotspots[loc]["delay"] += 1
            elif cat == "Maintenance Issue":
                hotspots[loc]["maintenance"] += 1

    locations = sorted(
        [
            {
                "location": loc,
                "complaint_count": d["count"],
                "safety_count": d["safety"],
                "delay_count": d["delay"],
                "maintenance_count": d["maintenance"],
            }
            for loc, d in hotspots.items()
        ],
        key=lambda x: -x["complaint_count"],
    )
    return {"total_hotspots": len(locations), "locations": locations}


# ══════════════════════════════════════════════════════════════════════
# EXPLAINABLE AI
# ══════════════════════════════════════════════════════════════════════

def get_ai_explanation(db: Session, complaint_id: str) -> dict:
    """Provides reasoning for AI severity and classification decisions."""
    complaint = db.query(Complaint).filter(Complaint.complaint_id == complaint_id).first()
    if not complaint:
        return {"error": f"Complaint '{complaint_id}' not found"}

    desc = (complaint.complaint_description or "").lower()
    reasons = []

    if complaint.severity == "Critical":
        if any(kw in desc for kw in ("accident", "crash", "injur", "hospitalise", "hospitalize")):
            reasons.append("Complaint mentions physical injury or accident — Critical severity applied.")
        elif any(kw in desc for kw in ("brake", "fail", "steering")):
            reasons.append("Complaint describes mechanical failure of a safety-critical component.")
        else:
            reasons.append("AI classified Critical severity due to keywords indicating immediate passenger welfare threat.")
    elif complaint.severity == "High":
        if any(kw in desc for kw in ("shout", "abuse", "rude", "assault", "threaten")):
            reasons.append("Complaint reports severe staff misconduct or aggressive behaviour.")
        else:
            reasons.append("High severity assigned due to significant service disruption or route skipping.")
    else:
        reasons.append(f"Severity classified as {complaint.severity or 'Medium'} — standard transit operational issue.")

    if complaint.priority_level in ("P1", "P2"):
        reasons.append(f"Priority {complaint.priority_level} escalated to guarantee rapid resolution for serious failure.")

    cats = _parse_json(complaint.ai_categories, [complaint.category] if complaint.category else [])
    reasons.append(f"Category assigned as {cats} via natural language content analysis.")

    return {
        "complaint_id": complaint_id,
        "severity": complaint.severity,
        "priority_level": complaint.priority_level,
        "assigned_department": complaint.assigned_department,
        "explanation": " ".join(reasons),
    }


# ══════════════════════════════════════════════════════════════════════
# REPORT DATA AGGREGATION
# ══════════════════════════════════════════════════════════════════════

def generate_report_data(db: Session, period: str) -> dict:
    """
    Aggregates all operational metrics needed for PDF/DOCX report generation.
    `period` is one of: "daily", "weekly", "monthly".
    """
    from datetime import timedelta

    now = datetime.now()
    if period == "daily":
        since = now - timedelta(days=1)
        period_label = f"Daily — {now.strftime('%d %b %Y')}"
    elif period == "weekly":
        since = now - timedelta(days=7)
        period_label = f"Weekly — Week ending {now.strftime('%d %b %Y')}"
    else:
        since = now - timedelta(days=30)
        period_label = f"Monthly — {now.strftime('%B %Y')}"

    # All complaints (for overall metrics)
    all_complaints = db.query(Complaint).all()
    # Period complaints
    period_complaints = [c for c in all_complaints if c.created_at and c.created_at >= since]

    total_all = len(all_complaints)
    total_period = len(period_complaints)

    resolved = sum(1 for c in all_complaints if c.complaint_status in ("Resolved", "Closed"))
    sla_breached = sum(1 for c in all_complaints if c.sla_status == "SLA Breached")
    sla_within = sum(1 for c in all_complaints if c.sla_status == "Within SLA")
    escalated = sum(1 for c in all_complaints if c.escalation_status)
    critical = sum(1 for c in all_complaints if c.severity == "Critical")

    resolution_rate = round((resolved / total_all * 100), 1) if total_all else 0
    sla_compliance = round((sla_within / max(sla_within + sla_breached, 1)) * 100, 1)

    # Category breakdown
    cat_counter: Counter = Counter(c.category for c in all_complaints if c.category)
    # Severity breakdown
    sev_counter: Counter = Counter(c.severity for c in all_complaints if c.severity)

    # Route risks
    try:
        routes = db.query(RouteRisk).order_by(RouteRisk.risk_score.desc()).limit(5).all()
        route_data = [
            {"route": r.route_number, "risk": r.risk_level, "score": round(r.risk_score, 2),
             "complaints": r.complaint_count}
            for r in routes
        ]
    except Exception:
        route_data = []

    # Driver risks
    try:
        drivers = db.query(DriverRisk).order_by(DriverRisk.risk_score.desc()).limit(5).all()
        driver_data = [
            {"driver": d.driver_identifier, "risk": d.risk_level, "score": round(d.risk_score, 2),
             "complaints": d.complaint_count}
            for d in drivers
        ]
    except Exception:
        driver_data = []

    # Bus risks
    try:
        buses = db.query(BusRisk).order_by(BusRisk.risk_score.desc()).limit(5).all()
        bus_data = [
            {"bus": b.bus_number, "maintenance_count": b.maintenance_count,
             "risk": b.maintenance_risk}
            for b in buses
        ]
    except Exception:
        bus_data = []

    # Governance recommendations (top 3)
    gov_recs = get_governance_recommendations(db)[:3]

    # Health index
    health = compute_health_index(db)

    return {
        "period": period.upper(),
        "period_label": period_label,
        "generated_at": now.strftime("%Y-%m-%d %H:%M:%S"),
        # Overall stats
        "total_complaints": total_all,
        "period_complaints": total_period,
        "resolved_count": resolved,
        "sla_breached_count": sla_breached,
        "sla_within_count": sla_within,
        "escalated_count": escalated,
        "critical_count": critical,
        "resolution_rate": resolution_rate,
        "sla_compliance": sla_compliance,
        # Breakdowns
        "categories": dict(cat_counter.most_common(8)),
        "severities": dict(sev_counter),
        # Risk tables
        "routes": route_data,
        "drivers": driver_data,
        "buses": bus_data,
        # Health & governance
        "health_score": health.get("transport_health_score", 0),
        "health_rating": health.get("rating", "N/A"),
        "health_metrics": health.get("metrics", {}),
        "governance": gov_recs,
    }


# ══════════════════════════════════════════════════════════════════════
# PDF REPORT BUILDER (ReportLab)
# ══════════════════════════════════════════════════════════════════════

def build_pdf_report(data: dict) -> bytes:
    """Builds a professional government-style PDF report and returns bytes."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable,
    )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50,
    )
    styles = getSampleStyleSheet()

    # ── Custom styles ──────────────────────────────────────────────────
    title_style = ParagraphStyle(
        "ReportTitle", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=20,
        textColor=colors.HexColor("#0f172a"), spaceAfter=4,
    )
    subtitle_style = ParagraphStyle(
        "ReportSubtitle", parent=styles["Normal"],
        fontName="Helvetica", fontSize=11,
        textColor=colors.HexColor("#475569"), spaceAfter=16,
    )
    h2_style = ParagraphStyle(
        "SectionH2", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=13,
        textColor=colors.HexColor("#1e40af"), spaceBefore=16, spaceAfter=8,
    )
    h3_style = ParagraphStyle(
        "SectionH3", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=11,
        textColor=colors.HexColor("#334155"), spaceBefore=10, spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "DocBody", parent=styles["Normal"],
        fontName="Helvetica", fontSize=10,
        textColor=colors.HexColor("#1e293b"), leading=15,
    )
    small_style = ParagraphStyle(
        "DocSmall", parent=styles["Normal"],
        fontName="Helvetica", fontSize=9,
        textColor=colors.HexColor("#64748b"),
    )

    # Colour palette
    HDR_BG = colors.HexColor("#1e293b")
    HDR_FG = colors.white
    ROW_ALT = colors.HexColor("#f8fafc")
    GRID_CLR = colors.HexColor("#e2e8f0")

    def _hdr_table_style():
        return TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), HDR_BG),
            ("TEXTCOLOR", (0, 0), (-1, 0), HDR_FG),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, ROW_ALT]),
            ("GRID", (0, 0), (-1, -1), 0.4, GRID_CLR),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ])

    story = []

    # ── Cover header ──────────────────────────────────────────────────
    story.append(Paragraph("Public Transport Authority", title_style))
    story.append(Paragraph(f"{data['period_label']} — Operations Report", subtitle_style))
    story.append(Paragraph(f"Generated: {data['generated_at']}", small_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#1e40af"), spaceAfter=14))

    # ── Section 1: Executive Summary ──────────────────────────────────
    story.append(Paragraph("1. Executive Summary", h2_style))
    health_color = (
        "#10b981" if data["health_score"] >= 70 else
        "#f59e0b" if data["health_score"] >= 50 else
        "#ef4444"
    )
    story.append(Paragraph(
        f"Transport Health Index: <font color='{health_color}'><b>{data['health_score']}/100 — {data['health_rating']}</b></font>",
        body_style,
    ))
    story.append(Spacer(1, 8))

    # KPI summary table
    kpi_data = [
        ["Metric", "Value"],
        ["Total Complaints (All Time)", str(data["total_complaints"])],
        [f"New Complaints ({data['period']} Period)", str(data["period_complaints"])],
        ["Resolved Complaints", str(data["resolved_count"])],
        ["Resolution Rate", f"{data['resolution_rate']}%"],
        ["SLA Compliance Rate", f"{data['sla_compliance']}%"],
        ["SLA Breaches", str(data["sla_breached_count"])],
        ["Escalated Complaints", str(data["escalated_count"])],
        ["Critical Severity Cases", str(data["critical_count"])],
    ]
    t = Table(kpi_data, colWidths=[270, 180])
    t.setStyle(_hdr_table_style())
    story.append(t)

    # ── Section 2: Complaints by Category ────────────────────────────
    story.append(Paragraph("2. Complaints by Category", h2_style))
    cat_rows = [["Category", "Count"]] + [[k, str(v)] for k, v in data["categories"].items()]
    if len(cat_rows) > 1:
        t2 = Table(cat_rows, colWidths=[270, 180])
        t2.setStyle(_hdr_table_style())
        story.append(t2)
    else:
        story.append(Paragraph("No category data available.", body_style))

    # ── Section 3: Complaints by Severity ────────────────────────────
    story.append(Paragraph("3. Complaints by Severity", h2_style))
    sev_rows = [["Severity", "Count"]] + [[k, str(v)] for k, v in data["severities"].items()]
    if len(sev_rows) > 1:
        t3 = Table(sev_rows, colWidths=[270, 180])
        t3.setStyle(_hdr_table_style())
        story.append(t3)
    else:
        story.append(Paragraph("No severity data available.", body_style))

    # ── Section 4: High Risk Routes ───────────────────────────────────
    story.append(Paragraph("4. High Risk Routes", h2_style))
    if data["routes"]:
        rt_rows = [["Route", "Risk Level", "Risk Score", "Complaints"]] + [
            [r["route"], r["risk"], f"{r['score']:.2f}", str(r["complaints"])]
            for r in data["routes"]
        ]
        t4 = Table(rt_rows, colWidths=[110, 110, 110, 120])
        t4.setStyle(_hdr_table_style())
        story.append(t4)
    else:
        story.append(Paragraph("No route risk data available. Run predictive analytics to generate.", body_style))

    # ── Section 5: Driver / Operator Risk ────────────────────────────
    story.append(Paragraph("5. Driver / Operator Risk Assessment", h2_style))
    if data["drivers"]:
        dr_rows = [["Bus (Proxy)", "Risk Level", "Risk Score", "Complaints"]] + [
            [d["driver"], d["risk"], f"{d['score']:.2f}", str(d["complaints"])]
            for d in data["drivers"]
        ]
        t5 = Table(dr_rows, colWidths=[110, 110, 110, 120])
        t5.setStyle(_hdr_table_style())
        story.append(t5)
    else:
        story.append(Paragraph("No driver risk data available.", body_style))

    # ── Section 6: Bus Health ─────────────────────────────────────────
    story.append(Paragraph("6. Bus Fleet Health", h2_style))
    if data["buses"]:
        bus_rows = [["Bus Number", "Maintenance Reports", "Maintenance Risk"]] + [
            [b["bus"], str(b["maintenance_count"]), b["risk"]]
            for b in data["buses"]
        ]
        t6 = Table(bus_rows, colWidths=[150, 150, 150])
        t6.setStyle(_hdr_table_style())
        story.append(t6)
    else:
        story.append(Paragraph("No bus risk data available.", body_style))

    # ── Section 7: Health Metrics ─────────────────────────────────────
    story.append(Paragraph("7. Health Sub-Index Metrics", h2_style))
    hm_rows = [["Dimension", "Score"]] + [
        [k.replace("_", " ").title(), str(v)]
        for k, v in data.get("health_metrics", {}).items()
    ]
    if len(hm_rows) > 1:
        t7 = Table(hm_rows, colWidths=[270, 180])
        t7.setStyle(_hdr_table_style())
        story.append(t7)

    # ── Section 8: Governance Recommendations ────────────────────────
    story.append(Paragraph("8. Strategic Governance Recommendations", h2_style))
    for i, rec in enumerate(data.get("governance", []), 1):
        story.append(Paragraph(f"{i}. {rec.get('title', 'Recommendation')}", h3_style))
        story.append(Paragraph(
            f"<b>Priority:</b> {rec.get('priority', '—')} &nbsp;|&nbsp; "
            f"<b>Category:</b> {rec.get('category', '—')}",
            body_style,
        ))
        story.append(Paragraph(rec.get("description", ""), body_style))
        story.append(Paragraph(f"<i>Estimated Impact: {rec.get('estimated_impact', '—')}</i>", small_style))
        story.append(Spacer(1, 6))

    # ── Footer ────────────────────────────────────────────────────────
    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=0.5, color=GRID_CLR))
    story.append(Paragraph(
        f"CONFIDENTIAL — Transport Complaint Agent · Generated {data['generated_at']}",
        small_style,
    ))

    doc.build(story)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes


# ══════════════════════════════════════════════════════════════════════
# DOCX REPORT BUILDER (python-docx)
# ══════════════════════════════════════════════════════════════════════

def build_docx_report(data: dict) -> bytes:
    """Builds a styled Word document and returns bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    def _set_cell_bg(cell, hex_color: str):
        """Apply background colour to a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _add_table(doc, headers: list[str], rows: list[list[str]], col_widths_in: list[float]):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        # Header row
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            _set_cell_bg(hdr_cells[i], "1E293B")
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
        # Data rows
        for row_idx, row in enumerate(rows):
            row_cells = table.add_row().cells
            for i, cell_val in enumerate(row):
                row_cells[i].text = str(cell_val)
                if row_idx % 2 == 1:
                    _set_cell_bg(row_cells[i], "F8FAFC")
                p = row_cells[i].paragraphs[0]
                if p.runs:
                    p.runs[0].font.size = Pt(9)
        # Column widths
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths_in):
                    cell.width = Inches(col_widths_in[i])
        return table

    # ── Cover ─────────────────────────────────────────────────────────
    title = doc.add_heading("Public Transport Authority", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    sub = doc.add_paragraph(f"{data['period_label']} — Operations Report")
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    gen = doc.add_paragraph(f"Generated: {data['generated_at']}")
    gen.runs[0].font.size = Pt(9)
    gen.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    doc.add_paragraph()

    # ── Section 1: Executive Summary ─────────────────────────────────
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        f"Transport Health Index: {data['health_score']}/100 — {data['health_rating']}"
    )
    _add_table(doc,
        ["Metric", "Value"],
        [
            ["Total Complaints (All Time)", str(data["total_complaints"])],
            [f"New Complaints ({data['period']} Period)", str(data["period_complaints"])],
            ["Resolved Complaints", str(data["resolved_count"])],
            ["Resolution Rate", f"{data['resolution_rate']}%"],
            ["SLA Compliance Rate", f"{data['sla_compliance']}%"],
            ["SLA Breaches", str(data["sla_breached_count"])],
            ["Escalated Complaints", str(data["escalated_count"])],
            ["Critical Severity Cases", str(data["critical_count"])],
        ],
        [3.5, 2.5],
    )
    doc.add_paragraph()

    # ── Section 2: Category Breakdown ────────────────────────────────
    doc.add_heading("2. Complaints by Category", level=1)
    if data["categories"]:
        _add_table(doc,
            ["Category", "Count"],
            [[k, str(v)] for k, v in data["categories"].items()],
            [3.5, 2.5],
        )
    else:
        doc.add_paragraph("No category data available.")
    doc.add_paragraph()

    # ── Section 3: Severity Breakdown ────────────────────────────────
    doc.add_heading("3. Complaints by Severity", level=1)
    if data["severities"]:
        _add_table(doc,
            ["Severity", "Count"],
            [[k, str(v)] for k, v in data["severities"].items()],
            [3.5, 2.5],
        )
    else:
        doc.add_paragraph("No severity data available.")
    doc.add_paragraph()

    # ── Section 4: Route Risks ────────────────────────────────────────
    doc.add_heading("4. High Risk Routes", level=1)
    if data["routes"]:
        _add_table(doc,
            ["Route", "Risk Level", "Risk Score", "Complaints"],
            [[r["route"], r["risk"], f"{r['score']:.2f}", str(r["complaints"])] for r in data["routes"]],
            [1.5, 1.5, 1.5, 1.5],
        )
    else:
        doc.add_paragraph("No route risk data. Run predictive analytics to generate.")
    doc.add_paragraph()

    # ── Section 5: Driver Risks ───────────────────────────────────────
    doc.add_heading("5. Driver / Operator Risk Assessment", level=1)
    if data["drivers"]:
        _add_table(doc,
            ["Bus (Proxy)", "Risk Level", "Risk Score", "Complaints"],
            [[d["driver"], d["risk"], f"{d['score']:.2f}", str(d["complaints"])] for d in data["drivers"]],
            [1.5, 1.5, 1.5, 1.5],
        )
    else:
        doc.add_paragraph("No driver risk data available.")
    doc.add_paragraph()

    # ── Section 6: Bus Fleet Health ───────────────────────────────────
    doc.add_heading("6. Bus Fleet Health", level=1)
    if data["buses"]:
        _add_table(doc,
            ["Bus Number", "Maintenance Reports", "Maintenance Risk"],
            [[b["bus"], str(b["maintenance_count"]), b["risk"]] for b in data["buses"]],
            [2.0, 2.0, 2.0],
        )
    else:
        doc.add_paragraph("No bus risk data available.")
    doc.add_paragraph()

    # ── Section 7: Health Metrics ─────────────────────────────────────
    doc.add_heading("7. Health Sub-Index Metrics", level=1)
    hm = data.get("health_metrics", {})
    if hm:
        _add_table(doc,
            ["Dimension", "Score"],
            [[k.replace("_", " ").title(), str(v)] for k, v in hm.items()],
            [3.5, 2.5],
        )
    doc.add_paragraph()

    # ── Section 8: Governance Recommendations ────────────────────────
    doc.add_heading("8. Strategic Governance Recommendations", level=1)
    for i, rec in enumerate(data.get("governance", []), 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(rec.get("title", "Recommendation"))
        run.bold = True
        run.font.size = Pt(11)
        doc.add_paragraph(
            f"Priority: {rec.get('priority', '—')}  |  Category: {rec.get('category', '—')}"
        ).runs[0].font.size = Pt(9)
        desc_p = doc.add_paragraph(rec.get("description", ""))
        desc_p.runs[0].font.size = Pt(10)
        imp = doc.add_paragraph(f"Estimated Impact: {rec.get('estimated_impact', '—')}")
        imp.runs[0].italic = True
        imp.runs[0].font.size = Pt(9)

    # ── Footer ────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        f"CONFIDENTIAL — Transport Complaint Agent · Generated {data['generated_at']}"
    )
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    return docx_bytes


# ══════════════════════════════════════════════════════════════════════
# POLICY SIMULATION (kept for backward compat)
# ══════════════════════════════════════════════════════════════════════

def simulate_policy(action_type: str, target: str) -> dict:
    """Simulates operational impact of policy decisions."""
    a = action_type.lower()
    if "add" in a or "bus" in a:
        return {"policy": f"Add buses to {target}", "predicted_complaint_reduction": "22%",
                "expected_wait_time_reduction": "15%", "impact_on_sla_compliance": "+8%",
                "cost_index": "Medium", "confidence": "High"}
    elif "weekend" in a or "frequency" in a:
        return {"policy": f"Increase weekend frequency by 20% on {target}",
                "predicted_complaint_reduction": "18%", "expected_wait_time_reduction": "12%",
                "impact_on_sla_compliance": "+5%", "cost_index": "Low-Medium", "confidence": "Medium-High"}
    elif "driver" in a or "train" in a:
        return {"policy": f"Conduct driver training for {target}",
                "predicted_complaint_reduction": "30%", "expected_wait_time_reduction": "0%",
                "impact_on_sla_compliance": "+3%", "cost_index": "Low", "confidence": "High"}
    return {"policy": f"Optimise operations on {target}", "predicted_complaint_reduction": "10%",
            "expected_wait_time_reduction": "8%", "impact_on_sla_compliance": "+4%",
            "cost_index": "Low", "confidence": "Medium"}
